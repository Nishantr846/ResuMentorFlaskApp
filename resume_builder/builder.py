from docx import Document
import tempfile
import os

def generate_resume_doc(data):
    template_name = data.get('template_name')
    template_path = None
    if template_name:
        template_path = os.path.join('static', 'template', template_name)
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
    else:
        doc = Document()
    doc.add_heading(data['name'], 0)
    doc.add_paragraph(f"Email: {data['email']}")
    doc.add_paragraph(f"Phone: {data['phone']}")
    doc.add_heading("Education", level=1)
    doc.add_paragraph(data['education'])
    doc.add_heading("Experience", level=1)
    doc.add_paragraph(data['experience'])
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(data['skills'])

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    return tmp.name